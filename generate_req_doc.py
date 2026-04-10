from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Helper functions ────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table_header_row(table, headers, bg='1F4E79'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(cell, bg)

def add_data_row(table, values, row_idx):
    row = table.add_row()
    bg = 'D6E4F0' if row_idx % 2 == 0 else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_bg(cell, bg)
    return row

def add_heading(doc, text, level, color_hex='1F4E79'):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
    return h

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

def add_paragraph(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )
    p.paragraph_format.space_after = Pt(4)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Page Margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ─── COVER PAGE ───────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('REQUIREMENT UNDERSTANDING DOCUMENT')
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('IT Total Cost of Ownership (TCO) &\nTechnology Business Management (TBM) Platform')
sub_run.bold = True
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.style = 'Table Grid'
meta_data = [
    ('Document Type',    'Requirement Understanding Document'),
    ('Source Document',  'IT_TCO_TBM_BRD.pdf'),
    ('Prepared By',      'Business Analysis Team'),
    ('Date Prepared',    datetime.date.today().strftime('%B %d, %Y')),
    ('Version',          '1.0'),
]
for i, (label, value) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    set_cell_bg(row.cells[0], 'D6E4F0')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ─── TABLE OF CONTENTS ───────────────────────────────────────────────────────
add_heading(doc, 'Table of Contents', 1)
toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Business Objectives'),
    ('3.', 'Project Scope'),
    ('4.', 'Phased Integration Strategy'),
    ('5.', 'Functional Requirements'),
    ('6.', 'Cost Allocation Model'),
    ('7.', 'Key Use Cases'),
    ('8.', 'Non-Functional Requirements'),
    ('9.', 'Expected Business Benefits'),
    ('10.', 'Assumptions & Constraints'),
    ('11.', 'Stakeholders & Roles'),
    ('12.', 'Success Criteria & KPIs'),
    ('13.', 'Open Questions & Risks'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r2 = p.add_run(item)
    r2.font.size = Pt(10)

doc.add_page_break()

# ─── SECTION 1: EXECUTIVE SUMMARY ────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', 1)
add_separator(doc)

add_paragraph(doc,
    'This Requirement Understanding Document is prepared based on the Business Requirement Document (BRD) '
    'for the IT Total Cost of Ownership (TCO) & Technology Business Management (TBM) Platform. '
    'The purpose of this document is to articulate a clear understanding of what is being built, '
    'why it is being built, and how it will be delivered.'
)

add_paragraph(doc,
    'The platform is designed to provide comprehensive visibility into IT costs across all services '
    'and applications within the enterprise. It aggregates cost data from multiple dimensions — '
    'People, Software, Hardware, and Cloud — and allocates those costs to individual applications '
    'and business units using consumption-based models.'
)

add_paragraph(doc,
    'The core value proposition of the platform is to transform IT from a cost center into a '
    'transparent, accountable, and strategically aligned function by enabling data-driven financial '
    'decision-making.'
)

doc.add_paragraph()

# Cost category table
add_paragraph(doc, 'Cost Categories Covered:', bold=True)
cat_table = doc.add_table(rows=1, cols=3)
cat_table.style = 'Table Grid'
add_table_header_row(cat_table, ['Cost Category', 'Examples', 'Data Source'])
categories = [
    ('People Cost',    'Employee salaries, contractor fees, consultant charges',      'HR / Financial Systems'),
    ('Software Cost',  'Application licenses, SaaS subscriptions, maintenance',       'OutSystems / ERP'),
    ('Hardware Cost',  'Servers, networking equipment, on-prem infrastructure',       'ServiceNow CMDB'),
    ('Cloud Cost',     'Azure compute, storage, networking, PaaS/SaaS consumption',  'Azure Cost Management API'),
]
for i, row_data in enumerate(categories):
    add_data_row(cat_table, row_data, i)
doc.add_paragraph()

# ─── SECTION 2: BUSINESS OBJECTIVES ──────────────────────────────────────────
add_heading(doc, '2. Business Objectives', 1)
add_separator(doc)
add_paragraph(doc, 'The platform is being built to achieve the following strategic business objectives:')

objectives = [
    ('Unified IT Cost View: ',           'Provide a single, consolidated view of IT costs across cloud, infrastructure, software, and personnel.'),
    ('Application-Level Transparency: ', 'Enable visibility into the complete cost profile of each enterprise application.'),
    ('Chargeback & Showback Models: ',   'Support both direct cost allocation (chargeback) and visibility without billing (showback) to business units.'),
    ('IT Financial Governance: ',        'Strengthen financial discipline across IT through structured cost reporting and accountability.'),
    ('Budget & Forecast Accuracy: ',     'Improve the quality of IT budgeting and financial forecasting using real-time and historical cost data.'),
    ('Cost Optimization: ',              'Proactively identify waste, redundant applications, and underutilized resources to reduce overall IT spend.'),
]
for bold_part, rest in objectives:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_paragraph()

# ─── SECTION 3: PROJECT SCOPE ─────────────────────────────────────────────────
add_heading(doc, '3. Project Scope', 1)
add_separator(doc)

add_heading(doc, '3.1 In-Scope', 2, color_hex='2E75B6')
in_scope = [
    'Ingestion of cloud cost data from Microsoft Azure via APIs',
    'Mapping of IT infrastructure and services to business applications using CMDB data',
    'Import of employee, contractor, and project cost data from internal financial systems',
    'Cost normalization, enrichment, and allocation engine',
    'Application TCO calculation using the defined cost allocation model',
    'Interactive dashboards for IT financial management and cost optimization',
    'Reporting by business unit, department, cost category, and service',
    'Role-based access control and secure authentication',
]
for item in in_scope:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '3.2 Out-of-Scope (Current Phase)', 2, color_hex='2E75B6')
out_scope = [
    'Multi-cloud integrations (AWS, GCP) — deferred to a future phase',
    'ERP and HR system integrations — deferred to a future phase',
    'Automated financial budgeting or forecasting engine',
    'Real-time streaming cost ingestion (batch ingestion is in scope)',
]
for item in out_scope:
    add_bullet(doc, item)

doc.add_paragraph()

# ─── SECTION 4: PHASED INTEGRATION STRATEGY ──────────────────────────────────
add_heading(doc, '4. Phased Integration Strategy', 1)
add_separator(doc)
add_paragraph(doc,
    'The platform will be delivered in phases, with each phase adding new data sources and capabilities. '
    'This phased approach reduces risk and allows for early delivery of value.'
)

phase_table = doc.add_table(rows=1, cols=4)
phase_table.style = 'Table Grid'
add_table_header_row(phase_table, ['Phase', 'Integration', 'Purpose', 'Status'])
phases = [
    ('Phase 1',       'Microsoft Azure Cost Management',   'Retrieve cloud consumption and cost data via Azure Cost Management APIs',                                   '🟢 Active'),
    ('Phase 2',       'ServiceNow CMDB',                   'Map infrastructure resources and services to business applications for cost attribution',                   '🔵 Planned'),
    ('Phase 3',       'OutSystems / Financial Systems',    'Import project costs, contracts, employee salary, and consultant cost data from internal financial systems', '🔵 Planned'),
    ('Future Phase',  'ERP / HR / Multi-cloud',            'Extend financial modeling with ERP, HR data, and integrate additional cloud providers (AWS, GCP)',           '⚪ Future'),
]
for i, row_data in enumerate(phases):
    add_data_row(phase_table, row_data, i)

doc.add_paragraph()

# ─── SECTION 5: FUNCTIONAL REQUIREMENTS ──────────────────────────────────────
add_heading(doc, '5. Functional Requirements', 1)
add_separator(doc)

func_reqs = [
    ('FR-001', 'Cloud Cost Ingestion',       'Ingest and normalize Azure cloud cost data via the Azure Cost Management API. Support daily or scheduled batch ingestion.',                                          'Must Have'),
    ('FR-002', 'CMDB Integration',           'Import configuration items (CIs), application records, and service mappings from ServiceNow CMDB to enable infrastructure-to-application cost attribution.',        'Must Have'),
    ('FR-003', 'People Cost Import',         'Import employee salary, contractor billing rates, and project time allocations from HR and financial systems to calculate people cost per application.',              'Must Have'),
    ('FR-004', 'Cost Allocation Engine',     'Allocate normalized costs to applications and business units using configurable allocation models based on consumption, headcount, and usage metrics.',               'Must Have'),
    ('FR-005', 'Application TCO Dashboard',  'Provide interactive dashboards showing the total cost of ownership per application, broken down by People, Software, Hardware, and Cloud cost categories.',          'Must Have'),
    ('FR-006', 'Business Unit Reporting',    'Generate reports showing IT cost allocation by business unit, department, and service line to support chargeback and showback models.',                              'Must Have'),
    ('FR-007', 'Cost Trend Analysis',        'Display historical cost trends at application, department, and platform level to support budget planning and forecasting.',                                          'Should Have'),
    ('FR-008', 'Optimization Insights',      'Surface recommendations and alerts for applications with high cost-per-user ratios, underutilized cloud resources, or redundant software licenses.',                 'Should Have'),
    ('FR-009', 'Cloud Spend Governance',     'Monitor and report Azure cloud spending patterns, budget thresholds, and anomaly detection.',                                                                        'Should Have'),
    ('FR-010', 'Role-Based Access Control',  'Implement RBAC to control visibility into cost data — e.g., business unit heads can only see their own cost allocations.',                                          'Must Have'),
]

fr_table = doc.add_table(rows=1, cols=4)
fr_table.style = 'Table Grid'
add_table_header_row(fr_table, ['Req. ID', 'Requirement', 'Description', 'Priority'])
for i, row_data in enumerate(func_reqs):
    r = add_data_row(fr_table, row_data, i)

doc.add_paragraph()

# ─── SECTION 6: COST ALLOCATION MODEL ────────────────────────────────────────
add_heading(doc, '6. Cost Allocation Model', 1)
add_separator(doc)

add_paragraph(doc, 'Core Formula:', bold=True)
formula_p = doc.add_paragraph()
formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula_run = formula_p.add_run('Application TCO  =  People Cost  +  Software Cost  +  Hardware Cost  +  Cloud Cost')
formula_run.bold = True
formula_run.font.size = Pt(13)
formula_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
doc.add_paragraph()

add_paragraph(doc, 'Allocation Rules:', bold=True)
alloc_rules = [
    ('People Cost Allocation: ',    'Allocated based on time-tracking data or headcount per application from project management and HR systems.'),
    ('Software Cost Allocation: ',  'Allocated based on license assignments or per-user software subscription costs mapped to applications.'),
    ('Hardware Cost Allocation: ',  'Allocated based on CMDB-mapped infrastructure resources (servers, storage, networking) associated with each application.'),
    ('Cloud Cost Allocation: ',     'Allocated directly from Azure Cost Management tags or subscription/resource group mappings to applications.'),
]
for bold_part, rest in alloc_rules:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_paragraph()

# ─── SECTION 7: KEY USE CASES ─────────────────────────────────────────────────
add_heading(doc, '7. Key Use Cases', 1)
add_separator(doc)

use_cases = [
    ('UC-001', 'Application Cost Transparency',
     'A CIO or IT leader can view the complete TCO of any enterprise application, broken down by cost category. This enables data-driven decisions on application rationalization and investment.',
     'CIO, IT Finance, Application Owners'),
    ('UC-002', 'Chargeback Model',
     'IT costs are allocated back to business units based on their actual consumption of IT services and applications. Business units are billed or shown their allocated IT spend monthly.',
     'Finance, Business Unit Heads'),
    ('UC-003', 'Showback Reporting',
     'Business units and departments receive visibility into how much IT resources they consume without being directly charged, improving awareness and accountability.',
     'Department Heads, IT Finance'),
    ('UC-004', 'Budget Planning',
     'Historical cost trends by application and business unit are used to build the annual IT budget. Finance and IT teams can model different scenarios based on forecast data.',
     'IT Finance, Budget Owners'),
    ('UC-005', 'Cost Optimization',
     'The platform identifies applications with high cost-per-user ratios, unused software licenses, and over-provisioned cloud resources, enabling targeted cost reduction initiatives.',
     'IT Operations, Cloud FinOps Team'),
    ('UC-006', 'Cloud Governance',
     'Cloud spending patterns across Azure subscriptions and resource groups are monitored. Anomalies, budget overruns, and untagged resources are surfaced for governance action.',
     'Cloud FinOps Team, IT Leadership'),
]

for uc in use_cases:
    uc_id, uc_name, uc_desc, uc_actors = uc
    add_heading(doc, f'{uc_id}: {uc_name}', 3, color_hex='2E75B6')
    add_paragraph(doc, uc_desc)
    p = doc.add_paragraph()
    r1 = p.add_run('Primary Actors: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(uc_actors)
    r2.font.size = Pt(10)
    doc.add_paragraph()

# ─── SECTION 8: NON-FUNCTIONAL REQUIREMENTS ──────────────────────────────────
add_heading(doc, '8. Non-Functional Requirements', 1)
add_separator(doc)

nfr_table = doc.add_table(rows=1, cols=3)
nfr_table.style = 'Table Grid'
add_table_header_row(nfr_table, ['Category', 'Requirement', 'Details'])
nfrs = [
    ('Scalability',      'SaaS Architecture',             'Platform must be built on a scalable SaaS architecture capable of handling growing data volumes from additional cloud providers and cost sources.'),
    ('Integration',      'API-Based Connectivity',        'All integrations with enterprise systems (Azure, CMDB, Financial Systems) must be API-based to ensure maintainability and flexibility.'),
    ('Security',         'Authentication & RBAC',         'Secure authentication (SSO/OAuth2) and role-based access control must be enforced. Sensitive cost data must be encrypted at rest and in transit.'),
    ('Performance',      'Analytics Dashboards',          'Dashboards must load within 3 seconds for standard queries. Large dataset queries should be supported with pagination and caching.'),
    ('Extensibility',    'Connectors & Plugins',          'The platform must support extensible connectors to easily add new data sources (e.g., AWS, GCP, ERP) without major re-architecture.'),
    ('Availability',     'High Availability',             'The platform should target 99.9% uptime with appropriate disaster recovery and failover mechanisms.'),
    ('Auditability',     'Audit Trails',                  'All cost allocations, data imports, and user actions should be logged for audit and compliance purposes.'),
]
for i, row_data in enumerate(nfrs):
    add_data_row(nfr_table, row_data, i)

doc.add_paragraph()

# ─── SECTION 9: EXPECTED BUSINESS BENEFITS ───────────────────────────────────
add_heading(doc, '9. Expected Business Benefits', 1)
add_separator(doc)

benefits = [
    ('Improved IT Financial Transparency: ',     'Stakeholders across IT and business units gain a clear, real-time view of where IT money is being spent.'),
    ('Better Budgeting & Forecasting: ',         'Historical cost data enables more accurate annual IT budgets and multi-year financial forecasts.'),
    ('Chargeback & Showback Enablement: ',       'Organizations can implement fair cost allocation models to drive accountability across business units.'),
    ('Application Portfolio Rationalization: ',  'Identification of redundant, underutilized, or high-cost applications enables informed rationalization decisions.'),
    ('Cloud Cost Optimization: ',                'Visibility into Azure spending patterns enables FinOps practices and targeted cloud cost reduction.'),
    ('Strategic IT Decision-Making: ',           'Leadership can make data-driven IT investment decisions using cost insights from the platform.'),
]
for bold_part, rest in benefits:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_paragraph()

# ─── SECTION 10: ASSUMPTIONS & CONSTRAINTS ───────────────────────────────────
add_heading(doc, '10. Assumptions & Constraints', 1)
add_separator(doc)

add_heading(doc, '10.1 Assumptions', 2, color_hex='2E75B6')
assumptions = [
    'Azure Cost Management APIs are accessible and the organization has the appropriate Azure subscriptions and permissions configured.',
    'ServiceNow CMDB data is reasonably accurate and contains application-to-infrastructure mappings that can be used for cost attribution.',
    'Employee cost data is available from HR or financial systems in a structured format (CSV, API, or database).',
    'Applications are uniquely identifiable in the CMDB and can be used as the primary key for cost allocation.',
    'Business units and departments have agreed-upon definitions and hierarchies for chargeback and showback reporting.',
    'The organization uses Azure as its primary cloud provider for Phase 1. Other cloud providers will be addressed in future phases.',
]
for item in assumptions:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, '10.2 Constraints', 2, color_hex='2E75B6')
constraints = [
    'Phase 1 is limited to Microsoft Azure only. AWS and GCP integrations are out of scope.',
    'Data refresh frequency will be batch-based (e.g., daily); real-time streaming is not in scope for initial phases.',
    'The platform must comply with the organization\'s data security and privacy policies.',
    'Integration with ERP/HR systems is deferred to Phase 3 and future phases.',
]
for item in constraints:
    add_bullet(doc, item)

doc.add_paragraph()

# ─── SECTION 11: STAKEHOLDERS & ROLES ────────────────────────────────────────
add_heading(doc, '11. Stakeholders & Roles', 1)
add_separator(doc)

stk_table = doc.add_table(rows=1, cols=3)
stk_table.style = 'Table Grid'
add_table_header_row(stk_table, ['Stakeholder', 'Role', 'Interest in Platform'])
stakeholders = [
    ('CIO / IT Leadership',       'Executive Sponsor',          'Strategic oversight of IT spend; portfolio rationalization decisions'),
    ('IT Finance Team',           'Platform Owner',             'Budget management, chargeback/showback models, cost reporting'),
    ('Cloud FinOps Team',         'Primary User',               'Cloud cost optimization, Azure governance, spending anomaly detection'),
    ('Application Owners',        'Consumer',                   'Visibility into total cost of their application; optimization guidance'),
    ('Business Unit Heads',       'Consumer',                   'Understanding of IT costs allocated to their department'),
    ('ServiceNow Admins',         'Integration Owner',          'Maintaining CMDB data accuracy and API access for integration'),
    ('Azure Platform Team',       'Integration Owner',          'Providing Cost Management API access and maintaining Azure tagging standards'),
    ('HR / Finance Systems Team', 'Integration Owner',          'Providing employee and contractor cost data feeds'),
    ('Development / Engineering', 'Builder',                    'Platform development, integration, and ongoing maintenance'),
]
for i, row_data in enumerate(stakeholders):
    add_data_row(stk_table, row_data, i)

doc.add_paragraph()

# ─── SECTION 12: SUCCESS CRITERIA & KPIs ─────────────────────────────────────
add_heading(doc, '12. Success Criteria & KPIs', 1)
add_separator(doc)

kpi_table = doc.add_table(rows=1, cols=3)
kpi_table.style = 'Table Grid'
add_table_header_row(kpi_table, ['KPI', 'Target', 'Measurement Method'])
kpis = [
    ('Azure cost data ingestion latency',       '< 24 hours (daily batch)',          'Monitoring of data pipeline execution times'),
    ('Application TCO coverage',                '> 90% of enterprise applications',  'Count of applications with complete TCO data vs. total'),
    ('Dashboard load time',                     '< 3 seconds for standard views',    'Performance monitoring and user testing'),
    ('Chargeback model accuracy',               '> 95% allocation accuracy',         'Reconciliation with finance team records'),
    ('User adoption rate',                      '> 80% of target users active',      'Platform login and usage analytics'),
    ('Cost optimization savings identified',    'Measurable savings opportunities',  'Cost reduction recommendations acted upon'),
    ('Data quality score (CMDB)',               '> 85% CI accuracy',                 'CMDB data validation and completeness reports'),
]
for i, row_data in enumerate(kpis):
    add_data_row(kpi_table, row_data, i)

doc.add_paragraph()

# ─── SECTION 13: OPEN QUESTIONS & RISKS ──────────────────────────────────────
add_heading(doc, '13. Open Questions & Risks', 1)
add_separator(doc)

add_heading(doc, '13.1 Open Questions', 2, color_hex='2E75B6')
open_qs = [
    'What is the Azure tagging strategy currently in place? Tags will be critical for accurate cloud cost attribution to applications.',
    'Is there an existing application portfolio / CMDB that is the authoritative source for application records?',
    'What is the preferred chargeback model — is it actual consumption-based or a flat allocation model?',
    'What are the agreed-upon cost allocation rules for shared services (e.g., networking, security tools used by multiple applications)?',
    'Which financial system (OutSystems, ERP) will be the authoritative source for employee and contractor cost data?',
    'What is the target user interface — a custom-built UI or integration with an existing BI tool (e.g., Power BI)?',
    'Are there regulatory or compliance requirements around how cost data is stored and accessed?',
]
for q in open_qs:
    add_bullet(doc, q)

doc.add_paragraph()
add_heading(doc, '13.2 Key Risks', 2, color_hex='2E75B6')
risks = [
    ('CMDB Data Quality Risk: ',       'If CMDB data is incomplete or inaccurate, cost attribution to applications will be unreliable. Mitigation: Data quality assessment and cleansing sprint before Phase 2.'),
    ('Azure Tagging Gaps: ',           'Untagged or inconsistently tagged Azure resources will result in unallocated cloud costs. Mitigation: Define and enforce an Azure tagging policy before Phase 1 go-live.'),
    ('Stakeholder Adoption Risk: ',    'Business units may resist chargeback models if cost allocations are not transparent or explainable. Mitigation: Involve stakeholders early; start with showback before chargeback.'),
    ('Integration Complexity Risk: ',  'Integrating with multiple enterprise systems (CMDB, Financial, HR) may introduce delays. Mitigation: Phased delivery reduces integration risk per phase.'),
    ('Data Privacy & Security Risk: ', 'Cost data containing salary and contractor rates must be handled securely. Mitigation: Implement RBAC and data masking for sensitive cost attributes.'),
]
for bold_part, rest in risks:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_paragraph()
add_separator(doc)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    f'Requirement Understanding Document  |  IT TCO & TBM Platform  |  v1.0  |  {datetime.date.today().strftime("%B %d, %Y")}'
)
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
footer_run.italic = True

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = 'Requirement understanding.doc'
doc.save(output_path)
print(f'Document saved: {output_path}')
